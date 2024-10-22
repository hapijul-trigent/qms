from fpdf import FPDF
from docx import Document
import io


def generate_pdf(data_list):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("helvetica", size=12)

    for i, data_dict in enumerate(data_list, start=1):
        # Adding a heading before each dictionary
        pdf.set_font("helvetica", 'B', size=14)
        pdf.cell(0, 10, f"Table {i}", ln=True, align='L')
        pdf.ln(5)  # Add some space

        # Add table headers
        pdf.set_font("helvetica", 'B', size=12)
        pdf.cell(90, 10, text="Checks", border=1)
        pdf.cell(90, 10, text="Status", border=1)
        pdf.ln()

        # Add table rows
        pdf.set_font("helvetica", size=12)
        for key, value in data_dict.items():
            pdf.cell(90, 10, text=str(key), border=1)
            pdf.cell(90, 10, text=str(value), border=1)
            pdf.ln()
        
        # Add a space between tables
        pdf.ln(10)

    return pdf


import streamlit as st
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

def generate_pdf(data_list):
    # Create a byte stream buffer
    buffer = BytesIO()
    
    # Create PDF document using the buffer
    pdf = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    heading_style = styles['Heading2']
    body_style = styles['BodyText']
    table_names = ('Checks', 'Label Information', 'Medicinal Ingredients')
    table_headers = (("Checks", "Status"), ('Key', 'Value'), ('Medicinal Ingredient', 'Quantity'))
    for i, data_dict in enumerate(data_list, start=0):
        
        elements.append(Paragraph(f"Table {table_names[i]}", heading_style))
        elements.append(Spacer(1, 12))

        table_data = [table_headers[i]]

        for key, value in data_dict.items():
            table_data.append([
                Paragraph(str(key), body_style), 
                Paragraph(str(value), body_style)
            ])
        
        
        table = Table(table_data, colWidths=[200, 300])

        
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))

        
        elements.append(table)
        elements.append(Spacer(1, 20))

    
    pdf.build(elements)
    buffer.seek(0)
    return buffer








def generate_docx(data_dict):
    doc = Document()
    doc.add_heading('Inspection Report', 0)

    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Checks'
    hdr_cells[1].text = 'Status'

    for key, value in data_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)

    return doc