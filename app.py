import streamlit as st
from fpdf import FPDF
from docx import Document
import io
from src.tools import load_yolo_model
from PIL import Image
import supervision as sv
import numpy as np
import cv2
from PIL import Image, ExifTags
import pandas as pd


favicon = Image.open("static/Canprev-Logo.png")
st.set_page_config(
    page_title="Canprev AI",
    page_icon=favicon,
    layout="wide",
    initial_sidebar_state="expanded"
)

st.image("static/CanPrev_4D-logo.jpg", width=800)
st.markdown("""
<style>

	.stTabs [data-baseweb="tab-list"] {
		gap: 3px;
    }

	.stTabs [data-baseweb="tab"] {
		height: 40px;
        white-space: pre-wrap;
		background-color: #13276F;
		border-radius: 4px 4px 0px 0px;
		gap: 1px;
		padding: 10px 2px 10px 2px;
        color: white;
    }

	.stTabs [aria-selected="true"] {
  		background-color: #FFFFFF;
        color: #13276F;
        border: 2px solid #13276F;
        border-bottom: none;
	}

</style>""", unsafe_allow_html=True)

# Load Model
model_side_QA = load_yolo_model('weights/model_side_view_qa.pt')
model_top_base_qa = load_yolo_model('weights/Top-Bottom-Checks-v2-40.pt')
model_unopened_botle_type_classification = load_yolo_model('weights/model_unopened_botle_type_classification.pt')
model_side_view_QA = {
    'dropper_bottle' : load_yolo_model('weights/model_dropper_bottle_side_view25.pt'),

}

model_powder_bottle_side_view = load_yolo_model('weights/model_powder_botle_side_view.pt')
model_liquid_bottle_side_view_80 = load_yolo_model('weights/model_liquid_botle_side_view_80.pt')
model_side_view_pill_bottle = load_yolo_model('weights/model_side_view_pill_botle.pt')


from collections import defaultdict
CHECKLIST = dict()
TOP_CHECKLIST, SIDE_CHECKLIST, BOTTOM_CHECKLIST = defaultdict(list), defaultdict(list), defaultdict(list)
SIDE_CHECKS = {'label', 'botle_with_neckband', 'curved_shoulder'}
SIDE_CHECKS_MAP = {'label': ('Label', 'Present'), 'botle_with_neckband': ('Neckband', 'Present'), 'curved_shoulder': ('Shoulder', 'Curved')}
SIDE_CHECKS_MAP_NEW = {
    'model_dropper_bottle_side_view_80': ['dropper_botle', 'dropper_botle_cap', 'dropper_botle_shoulder', 'label'],
    'model_powder_bottle_side_view': ['label', 'powder_botle', 'powder_botle_cap'],
    'model_liquid_bottle_side_view_80': ['Cytomatrix-Dermal-Liquid-Botle', 'Cytomatrix-Dermal-Liquid-Botle-Shoulder', 
            'Cytomatrix-Dermal-Liquid-Botle-With-Neckband', 'Magnesium_liquid_botle', 'Magnesium_liquid_botle_cap', 
            'Magnesium_liquid_botle_shoulder', 'label', 'Canprev-Gaba-Liquid-Botle', 
            'Canprev-Gaba-Liquid-Botle-Cap', 'Canprev-Omega-Liquid-Botle-Cap-With-Neckband', 
            'Canprev-Omega-Liquid-Botle', 'Canprev-Omega-Liquid-Botle-Shoulder', 'Canprev-Gaba-Liquid-Botle-shoulder'
        ],
}
TOP_CHECKS = {'Cap',}
TOP_CHECKS_MAP = {'Cap Type': 'Label Check', 'botle_with_neckband': 'Neckband Check', 'curved_shoulder': 'Shoulder Check'}
BOTTOM_CHECKS = {'Base',}
BOTTOM_CHECKS_MAP = {'Cap Type': 'Label Check', 'botle_with_neckband': 'Neckband Check', 'curved_shoulder': 'Shoulder Check'}


def update_CHECKLIST(key, value, CHECKLIST):
    CHECKLIST[key].append(value)


def identify_product_type(image, model):
    """Will check product type for Side view analysis"""
    try:
        pass
    except Exception as e:
        pass

def correct_image_orientation(image):
    """Correct the orientation of an image based on its EXIF data."""
    try:
        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation] == 'Orientation':
                break

        exif = image._getexif()
        if exif is not None:
            orientation = exif.get(orientation)
            if orientation == 3:
                image = image.rotate(180, expand=True)
            elif orientation == 6:
                image = image.rotate(270, expand=True)
            elif orientation == 8:
                image = image.rotate(90, expand=True)
    except (AttributeError, KeyError, IndexError):
        pass
    return image


def top_view_checks(image, model):
    
    result = model(image)[0]
    detections = sv.Detections.from_ultralytics(result)
    detections = detections[detections.confidence > .6]

    if len(detections.xyxy) == 0:
        for thing in TOP_CHECKS:
            update_CHECKLIST(thing, False, TOP_CHECKLIST)
    else:
        for thing in TOP_CHECKS:
            update_CHECKLIST(thing, detections.data['class_name'][0], TOP_CHECKLIST)


    return result.plot()



def side_view_checks(image, view_name, model):

    result = model(image)[0]
    detections = sv.Detections.from_ultralytics(result)
    detections = detections[detections.confidence > .6]
    
    if len(detections.xyxy) == 0:
        for thing in SIDE_CHECKS:
            update_CHECKLIST(thing, False, SIDE_CHECKLIST)
    else:
        for thing in SIDE_CHECKS:
            if thing in detections.data['class_name']:
                update_CHECKLIST(thing, True, SIDE_CHECKLIST)
            else:
                update_CHECKLIST(thing, False, SIDE_CHECKLIST)

    return result.plot()

def bottom_view_checks(image, model):
    
    result = model(image)[0]
    detections = sv.Detections.from_ultralytics(result)
    detections = detections[detections.confidence > .6]

    if len(detections.xyxy) == 0:
        for thing in BOTTOM_CHECKS:
            update_CHECKLIST(thing, False, BOTTOM_CHECKLIST)
    else:
        for thing in BOTTOM_CHECKS:
            update_CHECKLIST(thing, detections.data['class_name'][0], BOTTOM_CHECKLIST)


    return result.plot()


def simple_dict_to_streamlit_table(data_dict):
    """
    Converts a simple key-value dictionary into a stylish table using Streamlit.
    """
    global CHECKLIST
    print(data_dict['Cap'].split('-', 3))
    brand, content_type, _, cap_type = data_dict['Cap'].split('-', 3)
    data_dict.update({
        'Product Brand': brand,
        'Contains': content_type,
        'Cap Type': cap_type.replace('-Cap', '').replace('Cap', 'Plastic')
    })
    data_dict.pop('Cap')
    # _, _, data_dict['Base'] = data_dict['Base'].split('-', 2) if 'Canprev-Type1-' not in data_dict['Base'] else (None, None, data_dict['Base'].replace('Canprev-Type1-', ''))
    data_dict['Base'] = 'Good' if data_dict['Base'] else'Unknown'
    data_dict['Cap'] = 'Present' if data_dict['Cap Type'] else 'Unknown'
    data_dict['Shoulder Type'] = data_dict['Shoulder']
    data_dict['Shoulder'] = 'Good' if data_dict['Shoulder'] else 'Unknown'
    data_dict = {key: data_dict[key] for key in sorted(data_dict)}
    CHECKLIST = data_dict
    df = pd.DataFrame(list(data_dict.items()), columns=['Checks', 'Status'])
    st.dataframe(df, hide_index=True, use_container_width=True)



st.title("Unopened - Product Inspection Dashboard")
st.subheader("Upload Images")


col1, col2 = st.columns(2)
reset_enable = False
# if reset_enable:
#     if st.button('Clear'):
#         reset_enable = False
#         top_view_img, bottom_view_img, left_view_img, right_view_img, front_view_img, back_view_img = None, None, None, None, None, None
if not(reset_enable):
    with col1:
        top_view_img = st.file_uploader("Top View", type=["jpg", "png", "jpeg"])
        if top_view_img is not None:
            top_view_img = Image.open(top_view_img)
            top_view_img = correct_image_orientation(top_view_img)
    with col2:
        bottom_view_img = st.file_uploader("Bottom View", type=["jpg", "png", "jpeg"])
        if bottom_view_img is not None:
            bottom_view_img = Image.open(bottom_view_img)
            bottom_view_img = correct_image_orientation(bottom_view_img)

    col3, col4 = st.columns(2)
    with col3:
        left_view_img = st.file_uploader("Left View", type=["jpg", "png", "jpeg"])
        if left_view_img is not None:
            left_view_img = Image.open(left_view_img)
            left_view_img = correct_image_orientation(left_view_img)

    with col4:
        right_view_img = st.file_uploader("Right View", type=["jpg", "png", "jpeg"])
        if right_view_img is not None:
            right_view_img = Image.open(right_view_img)
            right_view_img = correct_image_orientation(right_view_img)



    col5, col6 = st.columns(2)
    with col5:
        front_view_img = st.file_uploader("Front View", type=["jpg", "png", "jpeg"])
        if front_view_img is not None:
            front_view_img = Image.open(front_view_img)
            front_view_img = correct_image_orientation(front_view_img)

    with col6:
        back_view_img = st.file_uploader("Back View", type=["jpg", "png", "jpeg"])
        if back_view_img is not None:
            back_view_img = Image.open(back_view_img)
            back_view_img = correct_image_orientation(back_view_img)



st.divider()

top, bottom, left, right, front, back = st.columns(6)
with top: top_view_panel = st.empty()
with bottom: bottom_view_panel = st.empty()
with left: left_view_panel = st.empty()
with right: right_view_panel = st.empty()
with front: front_view_panel = st.empty()
with back: back_view_panel = st.empty()

st.divider()


def display_update_checklist(results, view_name):
    global CHECKLIST
    if view_name == 'Side':
        cols = st.columns(4)
        nocol = 4
    else:
        cols = st.columns(2)
        nocol = 2
    
    idx = 0
    for key, value in results.items():
        
        col = cols[idx % nocol]
        actual_class = value[0]

        if view_name == 'Side':
            value = all(value)
            key = SIDE_CHECKS_MAP[key]
            # if value:
            #     col.checkbox(f"{key[0]}: {key[1]}", value=True, key=f"{view_name}_{key}")
            # else:
            #     col.checkbox(f"{key[0]}: {key[1]}", value=False, key=f"{view_name}_{key}")
            if value:
                CHECKLIST[key[0]] =  key[1]
        else:
            value = value[0]
            # if value:
            #     col.checkbox(f"{key}: {value}", value=True, key=f"{view_name}_{key}")
            # else:
            #     col.checkbox(f"{key}: {value}", value=False, key=f"{view_name}_{key}")
            CHECKLIST[key] =  value
        idx += 1


def merge_side_view_analysis(images, annotation_view_panels, model=None):
    model = model if model is not None else model_side_QA
    for view_name, image in images.items():
        if image:
            annotated_view = side_view_checks(image, view_name, model=model)
            annotation_view_panels[view_name].image(annotated_view, channels='bgr')
    return True




side_images = {
    "Left": left_view_img,
    "Right": right_view_img,
    "Front": front_view_img,
    "Back": back_view_img
}

annotation_view_panels = {
    'Top': top_view_panel,
    'Bottom': bottom_view_panel,
    "Left": left_view_panel,
    "Right": right_view_panel,
    "Front": front_view_panel,
    "Back": back_view_panel
}


if top_view_img and bottom_view_img and left_view_img and right_view_img and front_view_img and back_view_img:

    st.subheader("Unopened Bottle Checklist")
    top_check, bottom_check = st.columns(2)

    with top_check:
        top_annotated_view = top_view_checks(top_view_img, model=model_top_base_qa)
        annotation_view_panels['Top'].image(top_annotated_view, channels='bgr')
        display_update_checklist(TOP_CHECKLIST, "Top")

    with bottom_check:
        bottom_annotated_view = bottom_view_checks(bottom_view_img, model=model_top_base_qa)
        annotation_view_panels['Bottom'].image(bottom_annotated_view, channels='bgr')
        display_update_checklist(BOTTOM_CHECKLIST, "Bottom")

    product_type_results = model_unopened_botle_type_classification(front_view_img)[0]
    product_type = product_type_results.to_df().loc[0]['name']
    print(product_type)
    model = model_side_view_QA.get(product_type, None)
    CHECKLIST['Product Type'] = product_type.title().replace('_', ' ').replace('Botle', 'Bottle')

    merge_side_view_analysis(side_images, annotation_view_panels=annotation_view_panels, model=model)
    display_update_checklist(SIDE_CHECKLIST, "Side")

    reset_enable = True
    simple_dict_to_streamlit_table(CHECKLIST)









def generate_pdf(data_dict):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("helvetica", size=12)

    pdf.set_font("helvetica", 'B', size=12)
    pdf.cell(90, 10, text="Checks", border=1)
    pdf.cell(90, 10, text="Status", border=1)
    pdf.ln()


    pdf.set_font("helvetica", size=12)
    for key, value in data_dict.items():
        pdf.cell(90, 10, text=str(key), border=1)
        pdf.cell(90, 10, text=str(value), border=1)
        pdf.ln()

    return pdf


def generate_docx(data_dict):
    doc = Document()
    doc.add_heading('Inspection Report', 0)

    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Checks'
    hdr_cells[1].text = 'Status'

    for key, value in data_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)

    return doc




download_enabled = all([top_view_img, bottom_view_img, any(side_images.values())])

pdf_download_button, docs_download_button = st.columns([1,1], vertical_alignment='bottom')

with pdf_download_button:
    if download_enabled:
        pdf = generate_pdf(CHECKLIST)
        pdf_output = io.BytesIO()
        pdf.output(pdf_output)
        pdf_output.seek(0)
        st.download_button(label="Download PDF", data=pdf_output, file_name="inspection_report.pdf", mime="application/pdf")

    # with docs_download_button:
    #     if download_enabled:
    #         doc = generate_docx(CHECKLIST)
    #         doc_output = io.BytesIO()
    #         doc.save(doc_output)
    #         doc_output.seek(0)

    #         st.download_button(label="Download DOCX", data=doc_output, file_name="inspection_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

